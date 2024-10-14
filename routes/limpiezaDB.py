import os
import pandas as pd
from sqlalchemy import create_engine
from fastapi import HTTPException, FastAPI, APIRouter
from fastapi.responses import JSONResponse, FileResponse
from config.database import get_conexion, cursor, DB_HOST, DB_NAME, DB_PASSWORD, DB_PORT, DB_USER
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

app = FastAPI()
datos = APIRouter()

# Ruta donde se guardará el informe
OUTPUT_DIR = "/home/normas4b/Documentos/reportservice/reportsfile/administracion/empleados/"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "informe_tablas.docx")

# Verificar la conexión a la base de datos
if cursor is None:
    raise HTTPException(status_code=500, detail="No se pudo establecer la conexión a la base de datos")

def crear_motor():
    conexion = get_conexion()
    if conexion is None:
        raise HTTPException(status_code=500, detail="No se pudo establecer la conexión a la base de datos")

    db_url = f"postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}"
    motor = create_engine(db_url)
    return motor

def obtener_tablas(motor):
    query = """
    SELECT table_name 
    FROM information_schema.tables 
    WHERE table_schema = 'public';
    """
    tablas = pd.read_sql(query, motor)
    return tablas['table_name'].tolist()

@datos.post('/columnasNull/', tags=["Análisis"])
def revisar_columnas_null():
    motor = crear_motor()
    tablas = obtener_tablas(motor)

    resultado_columnas_null = []
    tablas_vacias = []

    for tabla in tablas:
        query_columnas = f"""
        SELECT column_name 
        FROM information_schema.columns 
        WHERE table_name = '{tabla}';
        """
        columnas = pd.read_sql(query_columnas, motor)
        columnas_vacias = []
        tabla_vacia = True

        query_vacia = f'SELECT COUNT(*) FROM "{tabla}";'
        total_filas = pd.read_sql(query_vacia, motor).iloc[0, 0]

        if total_filas == 0:
            tablas_vacias.append(tabla)
            continue  # Si está vacía, saltamos el análisis de NULLs.

        for columna in columnas['column_name']:
            if columna in ["deleted_at", "created_at", "updated_at"]:
                continue

            query_null = f"""
            SELECT COUNT(*) 
            FROM "{tabla}" 
            WHERE "{columna}" IS NULL;
            """
            count_nulls = pd.read_sql(query_null, motor).iloc[0, 0]

            if count_nulls > 0:
                columnas_vacias.append(columna)
                tabla_vacia = False

        if columnas_vacias:
            resultado_columnas_null.append({
                'tabla': tabla, 
                'columnasVacias': ', '.join(columnas_vacias)
            })

    # Generar el documento Word usando python-docx
    try:
        # Asegurar que el directorio de salida existe
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        # Crear un nuevo documento
        doc = Document()

        # Agregar un título centrado
        title = doc.add_heading('Informe de Columnas con Valores NULL y Tablas Vacías', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Agregar fecha de generación
        fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Fecha de generación: {fecha}\n")

        # Sección para Tablas con Columnas NULL
        if resultado_columnas_null:
            doc.add_heading('Tablas con Columnas con Valores NULL', level=1)
            for item in resultado_columnas_null:
                tabla = item['tabla']
                columnas = item['columnasVacias']
                doc.add_paragraph(f"Tabla: {tabla}", style='List Bullet')
                doc.add_paragraph(f"Columnas con NULLs: {columnas}", style='List Bullet 2')
        else:
            doc.add_heading('No se encontraron columnas con valores NULL en las tablas.', level=1)

        # Sección para Tablas Vacías
        if tablas_vacias:
            doc.add_heading('Tablas Vacías', level=1)
            for tabla in tablas_vacias:
                doc.add_paragraph(f"Tabla: {tabla}", style='List Bullet')
        else:
            doc.add_heading('No se encontraron tablas vacías.', level=1)

        # Guardar el documento
        doc.save(OUTPUT_FILE)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar el informe: {str(e)}")

    return JSONResponse({
        'tablasConNulls': resultado_columnas_null, 
        'tablasVacias': tablas_vacias,
        'mensaje': 'Informe generado con éxito. Use el endpoint /descargar-informe/ para obtener el documento.'
    })

# Endpoint adicional para descargar el informe
@datos.get('/descargar-informe/', tags=["Análisis"])
def descargar_informe():
    if not os.path.isfile(OUTPUT_FILE):
        raise HTTPException(status_code=404, detail="El informe no ha sido generado aún.")
    
    return FileResponse(
        path=OUTPUT_FILE, 
        filename=os.path.basename(OUTPUT_FILE), 
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

app.include_router(datos)
