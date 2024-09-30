# import pandas as pd
# from sqlalchemy import create_engine
# from fastapi import HTTPException, FastAPI, APIRouter
# from config.database import get_conexion, cursor, DB_HOST, DB_NAME, DB_PASSWORD, DB_PORT, DB_USER
# from docx import Document

# app = FastAPI()
# datos = APIRouter()

# if cursor is None:
#     raise HTTPException(status_code=500, detail="No se pudo establecer la conexión a la base de datos")

# def crear_motor():
#     conexion = get_conexion()
#     if conexion is None:
#         raise HTTPException(status_code=500, detail="No se pudo establecer la conexión a la base de datos")

#     db_url = f"postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}"
#     motor = create_engine(db_url)
#     return motor

# def obtener_tablas(motor):
#     query = """
#     SELECT table_name 
#     FROM information_schema.tables 
#     WHERE table_schema = 'public';
#     """
#     tablas = pd.read_sql(query, motor)
#     return tablas['table_name'].tolist()

# @datos.post('/columnasNull/', tags=["Análisis"])
# def revisar_columnas_null():
#     motor = crear_motor()
#     tablas = obtener_tablas(motor)

#     resultado_columnas_null = []
#     tablas_vacias = []

#     for tabla in tablas:
#         query_columnas = f"""
#         SELECT column_name 
#         FROM information_schema.columns 
#         WHERE table_name = '{tabla}';
#         """
#         columnas = pd.read_sql(query_columnas, motor)
#         columnas_vacias = []
#         tabla_vacia = True

#         query_vacia = f'SELECT COUNT(*) FROM "{tabla}";'
#         total_filas = pd.read_sql(query_vacia, motor).iloc[0, 0]

#         if total_filas == 0:
#             tablas_vacias.append(tabla)
#             continue  # Si está vacía, saltamos el análisis de NULLs.

#         for columna in columnas['column_name']:
#             if columna in ["deleted_at", "created_at", "updated_at"]:
#                 continue

#             query_null = f"""
#             SELECT COUNT(*) 
#             FROM "{tabla}" 
#             WHERE "{columna}" IS NULL;
#             """
#             count_nulls = pd.read_sql(query_null, motor).iloc[0, 0]

#             if count_nulls > 0:
#                 columnas_vacias.append(columna)
#                 tabla_vacia = False

#         if columnas_vacias:
#             resultado_columnas_null.append({
#                 'tabla': tabla, 
#                 'columnasVacias': ', '.join(columnas_vacias)
#             })

#     return {'tablasConNulls': resultado_columnas_null, 'tablasVacias': tablas_vacias}


# app.include_router(datos)


###########################

import pandas as pd
from sqlalchemy import create_engine
from fastapi import HTTPException, FastAPI, APIRouter
from config.database import get_conexion, cursor, DB_HOST, DB_NAME, DB_PASSWORD, DB_PORT, DB_USER
from docx import Document

app = FastAPI()
datos = APIRouter()

if cursor is None:
    raise HTTPException(status_code=500, detail="No se pudo establecer la conexión a la base de datos")

def crear_motor():
    conexion = get_conexion()
    if conexion is None:
        raise HTTPException(status_code=500, detail="No se pudo establecer la conexión a la base de datos")

    db_url = f"postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}"
    motor = create_engine(db_url)
    return motor

def obtener_tablas(motor):
    query = """
    SELECT table_name 
    FROM information_schema.tables 
    WHERE table_schema = 'public';
    """
    tablas = pd.read_sql(query, motor)
    return tablas['table_name'].tolist()

# Nueva función: detectar filas de prueba
def detectar_filas_prueba(motor, tablas):
    filas_prueba = []

    for tabla in tablas:
        query_columnas = f"""
        SELECT * 
        FROM "{tabla}" 
        WHERE 
            ({' OR '.join([f'"{columna}" ILIKE \'%test%\' OR "{columna}" ILIKE \'%prueba%\'' for columna in columnas['column_name']])})
            OR 
            ({' OR '.join([f'"{columna}"::TEXT > \'2100-01-01\' OR "{columna}"::TEXT < \'1900-01-01\'' for columna in columnas['column_name'] if 'date' in columna.lower()])})
            OR
            ({' OR '.join([f'"{columna}" ILIKE \'%@example.com%\' OR "{columna}" ILIKE \'%@test.com%\'' for columna in columnas['column_name'] if 'email' in columna.lower()])});
        """        
        columnas = pd.read_sql(query_columnas, motor)

        # Buscando registros con patrones comunes de pruebas
        query_prueba = """
                    SELECT * 
                    FROM "{}" 
                    WHERE 
                        ({})
                        OR 
                        ({})
                        OR
                        ({});
                    """.format(
                        tabla,
                        ' OR '.join([f'"{columna}" ILIKE \'%test%\' OR "{columna}" ILIKE \'%prueba%\'' for columna in columnas['column_name']]),
                        ' OR '.join([f'"{columna}"::TEXT > \'2100-01-01\' OR "{columna}"::TEXT < \'1900-01-01\'' for columna in columnas['column_name'] if 'date' in columna.lower()]),
                        ' OR '.join([f'"{columna}" ILIKE \'%@example.com%\' OR "{columna}" ILIKE \'%@test.com%\'' for columna in columnas['column_name'] if 'email' in columna.lower()])
)
        
        with motor.connect() as conexion:
            registros_prueba = pd.read_sql(query_prueba, conexion)

        if not registros_prueba.empty:
            filas_prueba.append({'tabla': tabla, 'registros': registros_prueba})

    return filas_prueba

# Modificada para incluir la detección de filas de prueba
@datos.post('/columnasNull/', tags=["Análisis"])
def revisar_columnas_null():
    motor = crear_motor()
    tablas = obtener_tablas(motor)

    resultado_columnas_null = []
    tablas_vacias = []

    for tabla in tablas:
        query_columnas = f"""
        SELECT column_name 
        FROM information_schema.columns 
        WHERE table_name = '{tabla}';
        """
        columnas = pd.read_sql(query_columnas, motor)
        columnas_vacias = []
        tabla_vacia = True

        query_vacia = f'SELECT COUNT(*) FROM "{tabla}";'
        total_filas = pd.read_sql(query_vacia, motor).iloc[0, 0]

        if total_filas == 0:
            tablas_vacias.append(tabla)
            continue  # Si está vacía, saltamos el análisis de NULLs.

        for columna in columnas['column_name']:
            if columna in ["deleted_at", "created_at", "updated_at"]:
                continue

            query_null = f"""
            SELECT COUNT(*) 
            FROM "{tabla}" 
            WHERE "{columna}" IS NULL;
            """
            count_nulls = pd.read_sql(query_null, motor).iloc[0, 0]

            if count_nulls > 0:
                columnas_vacias.append(columna)
                tabla_vacia = False

        if columnas_vacias:
            resultado_columnas_null.append({
                'tabla': tabla, 
                'columnasVacias': ', '.join(columnas_vacias)
            })

    # Detectar filas de prueba
    filas_prueba = detectar_filas_prueba(motor, tablas)

    return {
        'tablasConNulls': resultado_columnas_null, 
        'tablasVacias': tablas_vacias,
        'filasPrueba': filas_prueba
    }

# Nueva función: generar el documento Word
def generar_documento_word(tablas_con_nulls, tablas_vacias, filas_prueba):
    doc = Document()

    # Agregar título
    doc.add_heading('Reporte de Tablas con Valores NULL, Tablas Vacías y Filas de Prueba', 0)

    # Tablas con columnas NULL
    doc.add_heading('Tablas con Columnas NULL', level=1)
    if tablas_con_nulls:
        for tabla in tablas_con_nulls:
            doc.add_paragraph(f"Tabla: {tabla['tabla']}")
            doc.add_paragraph(f"Columnas con NULL: {tabla['columnasVacias']}")
    else:
        doc.add_paragraph("No se encontraron tablas con valores NULL.")

    # Tablas vacías
    doc.add_heading('Tablas Vacías', level=1)
    if tablas_vacias:
        for tabla in tablas_vacias:
            doc.add_paragraph(f"Tabla vacía: {tabla}")
    else:
        doc.add_paragraph("No se encontraron tablas vacías.")

    # Filas de prueba
    doc.add_heading('Filas de Prueba', level=1)
    if filas_prueba:
        for prueba in filas_prueba:
            doc.add_paragraph(f"Tabla: {prueba['tabla']}")
            for index, row in prueba['registros'].iterrows():
                doc.add_paragraph(f"Registro de prueba: {row.to_dict()}")
    else:
        doc.add_paragraph("No se encontraron filas de prueba.")

    # Guardar el documento
    doc.save("reporte_tablas_null_vacias_y_prueba.docx")

# Procesar datos al iniciar
@app.on_event("startup")
def procesar_datos_al_iniciar():
    conexion = get_conexion()

    if conexion is not None:
        motor = crear_motor()
        tablas = obtener_tablas(motor)

        # Revisar columnas con NULL y tablas vacías
        resultado = revisar_columnas_null()
        tablas_con_nulls = resultado['tablasConNulls']
        tablas_vacias = resultado['tablasVacias']
        filas_prueba = resultado['filasPrueba']

        # Generar el documento en Word
        generar_documento_word(tablas_con_nulls, tablas_vacias, filas_prueba)

        # Cerrar la conexión
        conexion.close()
    else:
        print("No se pudo establecer la conexión a la base de datos.")
